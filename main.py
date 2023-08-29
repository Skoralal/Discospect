import discord
from discord.ext import commands
from docx.api import Document
from docx.shared import Inches
# import urllib.request
import os
import os.path
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaFileUpload
from docx2pdf import convert
from bs4 import BeautifulSoup
import codecs
from homework import new_assignment

def init_lessons():
    global lessons
    lessons = []
    with codecs.open("C:\prog_questionmark\Calendar\html\\ver3.html", "r", "utf_8_sig") as html_file:
        content = html_file.read()
    soup = BeautifulSoup(content, "lxml")
    timetable = soup.find_all("table", class_="timetable")
    prep1 = soup.find_all("tr")
    # print(prep1)
    aboba = soup.select("[style*='white-space:pre-wrap']")
    for value in aboba:
        if value.text not in lessons:
            lessons.append(value.text)
    print(lessons)

# init_lessons()


SCOPES = ["https://www.googleapis.com/auth/drive"]

creds = None
folder_id_manual = "1Lz_tGNwv1v_cYHJH7JHfey3z17hDreSg"

if os.path.exists("token.json"):
    creds = Credentials.from_authorized_user_file("token.json", SCOPES)
# if not creds or not creds.valid:
    # if creds and creds.expired and creds.refresh_token:
    #     creds.refresh(Request())
    # else:
flow = InstalledAppFlow.from_client_secrets_file("C:\prog_questionmark\Discospect\secret\secret.json", SCOPES)
creds = flow.run_local_server(port = 0)

with open("token.json", "w") as token:
    token.write(creds.to_json())

try:
    service = build("drive", "v3", credentials=creds)

    response = service.files().list(q="name='3бЛ2' and mimeType='application/vnd.google-apps.folder'", spaces = "drive").execute()
    if not response["files"]:
        file_metadata = {
            "name": "3бЛ2",
            "mimeType": "application/vnd.google-apps.folder"
        }
        file = service.files().create(body=file_metadata, fiels="id").execute()
        folder_id = file.get("id")
    else:
        folder_id = response["files"][0]["id"]
        print(folder_id)
    


except HttpError as e:
    print("error ", str(e))


client = commands.Bot(command_prefix="!", intents = discord.Intents.all())

@client.command()
async def create(ctx, *, name=None):
    guild = ctx.message.guild
    for value in lessons:
        name = value
        if name == None:
            await ctx.send('Sorry, but you have to insert a name. Try again, but do it like this: `>create [channel name]`')
        else:
            await guild.create_text_channel(name)
            # await ctx.send(f"Created a channel named {name}") 

@client.event
async def on_ready():
    print("started")

@client.command()
async def hello(ctx):
    await ctx.send("aboba")

@client.command(pass_context = True)
async def get(ctx):
    text_channel_list = []
    alted = []
    for server in client.guilds:
        for channel in server.channels:
            if str(channel.type) == 'text':
                text_channel_list.append(channel)
    print("list", text_channel_list)
    for channel in text_channel_list:
        print(type(channel))
        text = ""
        try:
            document = Document(f"C:\prog_questionmark\Discospect\Папка с говной\{str(channel)}.docx")
        except:
            document = Document("C:\prog_questionmark\Discospect\Папка с говной\\blank.docx")
        async for message in channel.history():
            
            if message.content:
                
                if channel not in alted:
                    alted.append(channel)
                timestamp = message.created_at
                if message.content[0] == "$":
                    document.add_heading(f"{message.created_at}")
                    p = document.add_paragraph(message.content[1:])
                elif message.content[0] == "'":
                    prep = message.content[1:]
                    listed = prep.split("|")
                    print(listed)
                    aboba = new_assignment(listed[0], listed[1], channel, listed[2])
                    aboba.goog()
                elif message.content[0] == "&":
                    string = message.content[1:]
                    rows = string.split("\n")
                    data = []
                    for value in rows:
                        data.append(value.split("--"))
                    # print(data)
                    # print(len(rows))
                    # print(len(data[0]))
                    table = document.add_table(rows = len(rows), cols = len(data[0]))
                    # for id, name in data:
                    
                    #     # Adding a row and then adding data in it.
                    #     row = table.add_row().cells
                    #     # Converting id to string as table can only take string input
                    #     row[0].text = str(id)
                    #     row[1].text = name
                    for value in data:
                        row = table.add_row().cells
                        for i in range(len(value)):
                            row[i].text = value[i]
                else:
                    text = message.content + "\n" + text
                try:
                    pic_list = []
                    print(message.attachments)
                    for value in message.attachments:
                        await value.save(f"C:\prog_questionmark\Discospect\\temp_img_bin\{value.filename}")
                        pic_list.append(value.filename)
                except:
                    None
                
                print("eto", message.content)
                print(channel)
                print(message.created_at)
                await message.delete()

        p = document.add_paragraph(text)
        # space for tables

        #
        for value in pic_list:
            document.add_picture(f"C:\prog_questionmark\Discospect\\temp_img_bin\{value}", width=Inches(7))
        text = ""

        document.save(f"C:\prog_questionmark\Discospect\Папка с говной\{str(channel)}.docx")
    query = f"parents = '{folder_id_manual}'"
    response_q = service.files().list(q=query).execute()
    files = response_q.get("files")
    nextPageToken = response_q.get("nextPageToken")
    while nextPageToken:
        response_q = service.files().list(q=query, pageToken = nextPageToken).execute()
        files.extend(response_q.get("files"))
        nextPageToken = response_q.get("nextPageToken")

    print(files)
    
    print("alted", alted)
    for channel in alted:
        convert(f"C:\prog_questionmark\Discospect\Папка с говной\{str(channel)}.docx", f"C:\prog_questionmark\Discospect\PDF\{str(channel)}.pdf")
        file_metadata = {
            "name": f"{str(channel)}.pdf",
            "parents": [folder_id]
        }
        media = MediaFileUpload(f"C:\prog_questionmark\Discospect\PDF\{str(channel)}.pdf")
        
        try:
            fileid = [value["id"] for value in files if value["name"] == file_metadata["name"]]
            print(fileid)
            upload_file = service.files().update(media_body=media, fileId=fileid[0]).execute()
        except:
            upload_file = service.files().create(body=file_metadata, media_body=media, fields="id").execute()
    print("done")



client.run("MTEyNjY2NjkwMzU0NTQ1ODgxMQ.G_yZnx.Y6FVlNhV1Fe21buXOfRqIjl6CI-O8wQ8EQ563w")