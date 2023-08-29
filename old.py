import discord
from discord.ext import commands
from docx.api import Document
from docx.shared import Inches
# import urllib.request
import os
import os.path
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaFileUpload
from docx2pdf import convert


SCOPES = ["https://www.googleapis.com/auth/drive"]

creds = None

if os.path.exists("token.json"):
    creds = Credentials.from_authorized_user_file("token.json", SCOPES)
if not creds or not creds.valid:
    # if creds and creds.expired and creds.refresh_token:
    #     creds.refresh(Request())
    # else:
    flow = InstalledAppFlow.from_client_secrets_file("C:\prog_questionmark\Discospect\secret\secret.json", SCOPES)
    creds = flow.run_local_server(port = 0)
    
    with open("token.json", "w") as token:
        token.write(creds.to_json())

try:
    service = build("drive", "v3", credentials=creds)

    response = service.files().list(q="name='moscow_on_fire_1488' and mimeType='application/vnd.google-apps.folder'", spaces = "drive").execute()
    if not response["files"]:
        file_metadata = {
            "name": "moscow_on_fire_1488",
            "mimeType": "application/vnd.google-apps.folder"
        }
        file = service.files().create(body=file_metadata, fiels="id").execute()
        folder_id = file.get("id")
    else:
        folder_id = response["files"][0]["id"]
        print(folder_id)
    


except HttpError as e:
    print("error ", str(e))

#AIzaSyC7zppGkx6vS6ddhnDPj4f-fgWrZgPQOGo
# dog = Document(f"C:\prog_questionmark\Discospect\Папка с говной\основной.docx")

client = commands.Bot(command_prefix="!", intents = discord.Intents.all())


@client.event
async def on_ready():
    print("started")

@client.command()
async def hello(ctx):
    await ctx.send("aboba")

@client.command(pass_context = True)
async def get(ctx):
    text_channel_list = []
    alted = []
    for server in client.guilds:
        for channel in server.channels:
            if str(channel.type) == 'text':
                text_channel_list.append(channel)
    print(text_channel_list)
    for channel in text_channel_list:
        print(type(channel))
        text = ""
        async for message in channel.history():
            try:
                document = Document(f"C:\prog_questionmark\Discospect\Папка с говной\{str(channel)}.docx")
            except:
                document = Document("C:\prog_questionmark\Discospect\Папка с говной\\blank.docx")
            # text += message.content
            if message.content:
                text += " "
                if channel not in alted:
                    alted.append(channel)
            timestamp = message.created_at
            if message.content[0] == "$":
                document.add_heading(f"{message.created_at}")
                p = document.add_paragraph(message.content[1:])
            else:
                text += message.content
            try:
                print(message.attachments)
                # print(message.attachments[0].url)
                for value in message.attachments:
                    # pic = str(value.url)
                    # url = "http://site.meishij.net/r/58/25/3568808/a3568808_142682562777944.jpg"
                    # urllib.request.urlretrieve(pic, f"C:\prog_questionmark\Discospect\\temp_img_bin\{value.filename}")
                    await value.save(f"C:\prog_questionmark\Discospect\\temp_img_bin\{value.filename}")
                    document.add_picture(f"C:\prog_questionmark\Discospect\\temp_img_bin\{value.filename}", width=Inches(7))
            except:
                None
            
            print(message.content)
            print(channel)
            print(message.created_at)
            await message.delete()
        
        document.add_heading(f"{timestamp}")
        p = document.add_paragraph(text)

        document.save(f"C:\prog_questionmark\Discospect\Папка с говной\{str(channel)}.docx")
    print(alted)
    for channel in alted:
        convert(f"C:\prog_questionmark\Discospect\Папка с говной\{str(channel)}.docx", f"C:\prog_questionmark\Discospect\PDF\{str(channel)}.pdf")
        file_metadata = {
            "name": f"{str(channel)}.pdf",
            "parents": [folder_id]
        }
        media = MediaFileUpload(f"C:\prog_questionmark\Discospect\PDF\{str(channel)}.pdf")
        upload_file = service.files().create(body=file_metadata, media_body=media, fields="id").execute()



client.run("MTEyNjY2NjkwMzU0NTQ1ODgxMQ.G_yZnx.Y6FVlNhV1Fe21buXOfRqIjl6CI-O8wQ8EQ563w")